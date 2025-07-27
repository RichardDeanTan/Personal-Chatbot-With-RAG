import streamlit as st
import os
import re
import docx
from langchain_nvidia_ai_endpoints import ChatNVIDIA
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
import tempfile
import requests

st.set_page_config(
    page_title="RichBot - AI Personal Assistant",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded"
)

DEFAULT_DOC_PATH = "resource/PersonalProfile_RAG_purpose.docx"
TEMPLATE_URL = "https://raw.githubusercontent.com/RichardDeanTan/Personal-Chatbot-With-RAG/main/resource/Personal%20Profile%20-%20Template.docx"
EMBEDDING_MODEL = "sentence-transformers/paraphrase-multilingual-mpnet-base-v2"
PRIMARY_LLM_MODEL = "gotocompany/gemma-2-9b-cpt-sahabatai-instruct"

st.markdown("""
<style>
    .custom-chat-message {
        padding: 1rem;
        border-radius: 15px;
        margin-bottom: 1rem;
        max-width: 80%;
        word-wrap: break-word;
        color: white;
    }
    
    .custom-chat-user {
        background-color: #005C4B;
        margin-left: auto;
        margin-right: 0;
        text-align: right;
        border-bottom-right-radius: 5px;
    }
    
    .custom-chat-bot {
        background-color: #363636;
        margin-left: 0;
        margin-right: auto;
        text-align: left;
        border-bottom-left-radius: 5px;
    }
    
    .chat-container {
        display: flex;
        flex-direction: column;
        gap: 1rem;
    }
    
    .user-message-container {
        display: flex;
        justify-content: flex-end;
        width: 100%;
    }
    
    .bot-message-container {
        display: flex;
        justify-content: flex-start;
        width: 100%;
    }

    div[data-testid="stButton"] > button,
    div[data-testid="stDownloadButton"] > button,
    div[data-testid="stSidebarButton"] > button {
        width: 100% !important;
        border-radius: 20px;
        border: none;
        background: linear-gradient(90deg, #00D4AA, #0084FF);
        color: white;
        transition: background 0.3s ease;
    }

    div[data-testid="stButton"] > button:hover,
    div[data-testid="stDownloadButton"] > button:hover,
    div[data-testid="stSidebarButton"] > button:hover {
        background: linear-gradient(90deg, #00B391, #006EDC);
        color: white;
    }
</style>
""", unsafe_allow_html=True)

# === INITIALIZATION ===
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

if 'retriever' not in st.session_state:
    st.session_state.retriever = None

if 'llm' not in st.session_state:
    st.session_state.llm = None

if 'document_processed' not in st.session_state:
    st.session_state.document_processed = False

if 'current_doc_name' not in st.session_state:
    st.session_state.current_doc_name = "Default Richard's Profile"

@st.cache_resource
def load_api_key():
    try:
        nvidia_api_key = st.secrets["NVIDIA_API_KEY"]
        os.environ["NVIDIA_API_KEY"] = nvidia_api_key
        return True
    except Exception as e:
        st.error(f"❌ Failed to load API key: {str(e)}")
        return False

def load_document(file_path=None, uploaded_file=None):
    try:
        if uploaded_file is not None:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_file_path = tmp_file.name
            
            doc = docx.Document(tmp_file_path)
            os.unlink(tmp_file_path) # Clean temp file
        else:
            # Handle file path
            doc = docx.Document(file_path)
        
        full_text = []
        
        for para in doc.paragraphs:
            # Check list items
            if para._p.pPr and para._p.pPr.numPr:
                indent_level = para._p.pPr.numPr.ilvl.val
                indentation = '    ' * indent_level

                if indent_level == 0:
                    bullet = '•'
                elif indent_level == 1:
                    bullet = 'o'
                else:
                    bullet = '-'
                
                full_text.append(f"{indentation}{bullet} {para.text}")
            else:
                full_text.append(para.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Error loading document: {str(e)}")

def create_logical_chunks(text_content):
    pattern = r'\n(?=\d+\.\s[A-Z])'
    chunks = re.split(pattern, text_content)
    
    cleaned_chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
    
    if len(cleaned_chunks) > 1 and not cleaned_chunks[0].strip().startswith('1.'):
        cleaned_chunks[1] = cleaned_chunks[0] + '\n\n' + cleaned_chunks[1]
        cleaned_chunks.pop(0)

    return cleaned_chunks

@st.cache_resource
def create_vector_store(chunks):
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
    vector_store = FAISS.from_texts(texts=chunks, embedding=embeddings)
    return vector_store.as_retriever(search_kwargs={'k': 5})

def create_llm(temperature=0.7, top_p=0.7, max_tokens=256):
    return ChatNVIDIA(
        model=PRIMARY_LLM_MODEL,
        temperature=temperature,
        top_p=top_p,
        max_tokens=max_tokens,
    )

def create_prompt_template():
    template = """
Anda adalah 'RichBot', sebuah asisten AI dengan kepribadian yang santai, ramah, dan informatif. Anggap diri Anda seperti seorang teman yang sedang bersemangat menceritakan profil Richard Dean Tanjaya kepada pengguna.

**--- GAYA BAHASA & PERSONALITAS ---**
1.  **Santai & Tidak Kaku:** Gunakan bahasa sehari-hari yang sopan. Hindari jawaban yang terlalu formal atau terdengar seperti skrip. Boleh menggunakan kata-kata seperti "nih", "lho", "sih", "keren, kan?", atau "hehe" jika konteksnya pas.
2.  **Variasi Awalan Kalimat:** JANGAN selalu memulai jawaban dengan "Tentu!". Coba variasikan dengan "Oh, kalau soal itu...", "Oke, jadi gini...", "Siap! Untuk...", "Wah, pertanyaan bagus!", atau langsung ke poin utama.
3.  **Proaktif & Membantu:** Jika pengguna tampak bingung atau bertanya secara umum, bantu arahkan percakapan dengan menawarkan beberapa topik menarik.

**--- ATURAN UTAMA (SANGAT PENTING) ---**
1. **SELALU PERIKSA CONTEXT TERLEBIH DAHULU:** Sebelum menjawab apapun, WAJIB baca dan analisis CONTEXT yang diberikan.
2. **JANGAN PERNAH MENGARANG:** Hanya gunakan informasi yang ADA di CONTEXT. Jika tidak ada, katakan dengan jelas.
3. **KOREKSI INFORMASI YANG SALAH:** Jika pengguna menyatakan sesuatu yang BERTENTANGAN dengan CONTEXT, WAJIB koreksi dengan tegas tapi ramah.
4. **BAHASA INDONESIA 100%:** Selalu balas dalam Bahasa Indonesia, apapun bahasa pertanyaan pengguna.
5. **JANGAN SEBUT KATA "CONTEXT":** Jika informasi tidak tersedia, gunakan frasa seperti "di informasi yang aku punya", "dari yang aku tau", "berdasarkan profil Richard", atau "di data yang tersedia".

**ATURAN PALING PENTING: KOREKSI PENGGUNA JIKA SALAH**
Jika pertanyaan atau pernyataan pengguna SALAH atau BERTENTANGAN dengan `CONTEXT`, tugas utama Anda adalah MENGOREKSI mereka dengan ramah dan tegas. Jangan pernah setuju dengan informasi yang salah.

**--- SKENARIO & CONTOH WAJIB ---**

**Skenario 1: Pertanyaan Topik Baru tentang Richard**
-   *User Question:* "apa saja proyeknya?"
-   *Jawaban Ideal Anda:* "Oh, soal proyek ya? Ada beberapa yang keren nih di portofolionya, seperti Personal AI Chatbot, Prediksi Obesitas, dan Analisis Sentimen Saham. Mau aku ceritain lebih dalam soal salah satunya?"

**Skenario 2: Pertanyaan Meta ("Kamu bisa apa?")**
-   *User Question:* "kamu bisa kasi tau aku apa aja?"
-   *Jawaban Ideal Anda:* "Aku bisa ceritain banyak hal tentang Richard. Mulai dari Informasi Pribadi, Pendidikan, Keterampilan (Skills), Pengalaman kerja, Proyek-proyeknya, sampai Sertifikasi yang dia punya. Kita mulai dari mana enaknya?"

**Skenario 3: Pertanyaan Navigasi ("Bahas apa lagi?")**
-   *User Question:* "next kita mau bahas apa?" atau "apalagi yang menarik?"
-   *Jawaban Ideal Anda:* "Masih banyak lho yang bisa kita bahas! Ada topik soal Pengalaman kerja, Latar Belakang Pendidikan, atau Keterampilan teknisnya. Kamu tertarik sama yang mana?"

**Skenario 4: Pertanyaan Follow-up untuk Detail**
-   *Chat History:*
    RichBot: "...Mau aku ceritain lebih dalam soal salah satunya?"
-   *New User Question:* "jelaskan yang chatbot"
-   *Jawaban Ideal Anda:* "Oke, jadi gini. Proyek Personal AI Chatbot itu intinya dia bikin chatbot pribadi pakai teknologi AI dan RAG. Tujuannya supaya chatbotnya bisa ngasih jawaban yang pas berdasarkan data pribadinya dia. Keren, kan?" *(Setelah ini, berhenti dan tunggu respons pengguna).*

**Skenario 5: Menangani Respon Singkat/Pujian ("menarik", "hmm")**
-   *New User Question:* "menarik" atau "hmm oke"
-   *Jawaban Ideal Anda:* "Asiik, kalau kamu tertarik! Mau lanjut liat bagian lain dari profilnya, atau ada pertanyaan spesifik mungkin?"

**Skenario 6: Jawaban TIDAK ADA di Konteks**
-   *User Question:* "Dia bisa main musik gak?"
-   *Jawaban Ideal Anda:* "Wah, kalau soal kemampuan main musik, sepertinya belum ada infonya nih di profilnya. Mungkin bisa jadi hobi baru, hehe."

---
**CONTEXT:**
{context}

**CHAT HISTORY (percakapan terakhir):**
{chat_history}

**User Question:** {question}
**Jawaban Anda (dalam Bahasa Indonesia yang santai):**
"""
    return PromptTemplate(
        template=template,
        input_variables=["context", "chat_history", "question"]
    )

def format_chat_history(history):
    if not history:
        return "No conversation yet."
    return "\n".join([f"User: {turn['user']}\nRichBot: {turn['bot']}" for turn in history[-1:]])  # Only last turn

def display_chat_message(message, is_user=False):
    if is_user:
        st.markdown(f"""
        <div class="user-message-container">
            <div class="custom-chat-message custom-chat-user">
                👤 <strong>You:</strong><br>{message}
            </div>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <div class="bot-message-container">
            <div class="custom-chat-message custom-chat-bot">
                👾<strong>RichBot:</strong><br>{message}
            </div>
        </div>
        """, unsafe_allow_html=True)

def process_document(uploaded_file=None):
    try:
        with st.spinner("🔄 Processing document..."):
            if uploaded_file:
                document_text = load_document(uploaded_file=uploaded_file)
                st.session_state.current_doc_name = uploaded_file.name
            else:
                document_text = load_document(DEFAULT_DOC_PATH)
                st.session_state.current_doc_name = "Default Richard's Profile"
            
            logical_chunks = create_logical_chunks(document_text)
            st.session_state.retriever = create_vector_store(logical_chunks)
            st.session_state.document_processed = True
            
            return len(logical_chunks)
    except Exception as e:
        st.error(f"❌ Error processing document: {str(e)}")
        return None

def refresh_chat():
    st.session_state.chat_history = []
    st.rerun()

# === SIDEBAR ===
with st.sidebar:
    if st.button("🔄 Start a New Chat", key="refresh_chat", help="Start a new conversation"):
        refresh_chat()
    
    st.markdown("---")
    
    st.header("About RichBot")
    st.markdown("""
    **RichBot** adalah AI chatbot yang dirancang khusus untuk memperkenalkan profil Richard Dean Tanjaya. 
    
    Dengan teknologi RAG (Retrieval-Augmented Generation), RichBot dapat memberikan informasi yang akurat tentang:
    - Informasi Pribadi
    - Pendidikan  
    - Pengalaman Kerja
    - Proyek-proyek
    - Sertifikasi
    - Keterampilan Teknis
    """)

    st.info("""
    Teknologi yang Digunakan:
    - 🔍 **RAG**  
    Menggunakan **FAISS** sebagai vector store untuk mencari potongan dokumen yang relevan berdasarkan pertanyaan pengguna.
    - 🧠 **Session-based Memory**  
    Menyimpan riwayat percakapan secara manual menggunakan `st.session_state` (hanya **1 interaksi** terakhir yang digunakan sebagai konteks).
    """)

    
    st.markdown("---")
    
    st.header("📄 Document Management")
    st.info(f"Current document: **{st.session_state.current_doc_name}**")    
    st.warning(
        "**⚠️ Important:**\n"
        "Dokumennya harus mengikuti format template untuk hasil optimal. "
        "Gunakan penomoran pada setiap bagian (1., 2., 3., dst.) dengan judul yang jelas."
    )

    try:
        response = requests.get(TEMPLATE_URL)
        if response.status_code == 200:
            st.download_button(
                label="📥 Download Template",
                data=response.content,
                file_name="Document_Template.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("❌ Failed to fetch template from GitHub.")
    except FileNotFoundError:
        st.error("Template file not found.")
    
    uploaded_file = st.file_uploader(
        "Upload your own profile document",
        type=['docx'],
        help="Upload a .docx file to replace the default profile"
    )

    # Menentukan file new, unprocessed file has been uploaded
    is_new_file_uploaded = (uploaded_file is not None) and (uploaded_file.name != st.session_state.current_doc_name)

    if is_new_file_uploaded:
        if st.button("🔄 Process New Document"):
            chunks_count = process_document(uploaded_file=uploaded_file)
            if chunks_count:
                st.success(f"✅ Document '{uploaded_file.name}' processed! Created {chunks_count} chunks.")
                st.balloons()
                st.rerun()
    else:
        # No new file is pending
        if st.button("📖 Load Default Document"):
            if st.session_state.current_doc_name == "Default Richard's Profile":
                st.info("Default document is already loaded.")
            else:
                chunks_count = process_document()
                if chunks_count:
                    st.success(f"✅ Default document reloaded! Created {chunks_count} chunks.")
                    st.rerun()

    st.markdown("---")
    
    st.header("⚙️ Model Parameters")
    temperature = st.slider(
        "Temperature",
        min_value=0.1,
        max_value=1.0,
        value=0.7,
        step=0.1,
        help="Mengontrol tingkat randomness dalam respons. Semakin tinggi = semakin kreatif, semakin rendah = semakin terfokus."
    )
    
    top_p = st.slider(
        "Top P",
        min_value=0.1,
        max_value=1.0,
        value=0.7,
        step=0.1,
        help="Mengontrol keragaman respons melalui metode nucleus sampling."
    )
    
    max_tokens = st.slider(
        "Max Tokens",
        min_value=50,
        max_value=512,
        value=256,
        step=50,
        help="Maximum length dari response"
    )
    
    vector_k = st.slider(
        "Document Retrieval (K)",
        min_value=1,
        max_value=10,
        value=5,
        step=1,
        help="Jumlah potongan dokumen relevan yang akan diambil."
    )

def main():
    # === MAIN INTERFACE ===
    st.title("RichBot - Personal AI Assistant")
    st.markdown("*Selamat datang! Saya siap membantu Anda mengenal Richard Dean Tanjaya.*")

    try:
        load_api_key()
    except ValueError as e:
        st.error(f"❌ API Key Error: {e}")
        st.stop()

    # Auto process default document
    if not st.session_state.document_processed and not st.session_state.retriever:
        with st.spinner("Loading default document..."):
            chunks_count = process_document()
            if chunks_count:
                st.success(f"✅ Default document loaded with {chunks_count} chunks!")

    # Inisialisasi LLM + Default Parameter
    if st.session_state.retriever:
        st.session_state.llm = create_llm(temperature, top_p, max_tokens)    
        st.session_state.retriever.search_kwargs['k'] = vector_k

    # Display chat history
    st.markdown('<div class="chat-container">', unsafe_allow_html=True)

    if not st.session_state.chat_history and st.session_state.document_processed:
        # Initial greeting
        display_chat_message("Halo, perkenalkan namaku RichBot! Aku adalah AI Chatbot yang siap membantumu mengenal Richard. Silakan ajukan pertanyaanmu.", is_user=False)

    for message in st.session_state.chat_history:
        display_chat_message(message["user"], is_user=True)
        display_chat_message(message["bot"], is_user=False)

    st.markdown('</div>', unsafe_allow_html=True)

    # Chat input
    if prompt := st.chat_input("Ketik pertanyaan Anda di sini..."):
        if not st.session_state.retriever or not st.session_state.llm:
            st.error("❌ Please process a document first!")
            st.stop()
        
        # Add user message to history and display
        st.session_state.chat_history.append({"user": prompt, "bot": ""})
        st.rerun() # Rerun to display message user

    if st.session_state.chat_history and st.session_state.chat_history[-1]["bot"] == "":
        last_message = st.session_state.chat_history[-1]
        prompt = last_message["user"]

        # Generate Chatbot response
        with st.spinner("RichBot is thinking..."):
            try:
                # RAG
                retrieved_docs = st.session_state.retriever.invoke(prompt)
                context_text = "\n\n---\n\n".join([doc.page_content for doc in retrieved_docs])
                # Prompt
                prompt_template = create_prompt_template()
                # Memory
                recent_history = st.session_state.chat_history[:-1]
                
                formatted_prompt = prompt_template.format(
                    context=context_text,
                    chat_history=format_chat_history(recent_history),
                    question=prompt
                )
                
                # Generate response
                response = st.session_state.llm.invoke(formatted_prompt)
                full_response = response.content.strip()

                # Update message terakhir dengan Chatbot's response
                st.session_state.chat_history[-1]["bot"] = full_response
                
                # Rerun to display bot's response
                st.rerun()
                
            except Exception as e:
                st.error(f"❌ Error generating response: {str(e)}")

    # Footer
    st.markdown("---")
    st.markdown(
        """
        <div style="text-align: center; color: #666; font-size: 0.8em;">
            💡 Tip: Use the sidebar to customize model parameters and upload your own documents!
        </div>
        """, 
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()